"""
Document processor for medical documents including PDFs, drug databases, and clinical protocols.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
import fitz  # PyMuPDF
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

from .config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class MedicalDocumentProcessor:
    """Process medical documents and extract text for RAG pipeline."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,  # Increased from 1000 to 2000
            chunk_overlap=400,  # Increased from 200 to 400
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]  # Better separators for medical text
        )
        self.supported_formats = settings.supported_formats
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """
        Process a medical document and return chunks.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of document chunks
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            file_extension = file_path.suffix.lower().lstrip('.')
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract text based on file type
            if file_extension == 'pdf':
                text = self._extract_pdf_text(file_path)
            elif file_extension == 'docx':
                text = self._extract_docx_text(file_path)
            elif file_extension == 'txt':
                text = self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create Langchain documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "source": str(file_path),
                        "chunk_id": i,
                        "file_type": file_extension,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
            
            logger.info(f"Processed {file_path}: {len(documents)} chunks created")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods for better coverage."""
        text = ""
        
        try:
            # Method 1: PyMuPDF (fitz) - better for complex PDFs
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
            
            # If PyMuPDF didn't extract much text, try pdfplumber
            if len(text.strip()) < 100:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            # Fallback to PyPDF2
            if len(text.strip()) < 100:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                        
        except Exception as e:
            logger.warning(f"Error with PyMuPDF/pdfplumber, trying PyPDF2: {e}")
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e2:
                logger.error(f"All PDF extraction methods failed: {e2}")
                raise
        
        return text.strip()
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            raise
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of all document chunks
        """
        directory = Path(directory_path)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        all_documents = []
        
        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower().lstrip('.') in self.supported_formats:
                try:
                    documents = self.process_document(str(file_path))
                    all_documents.extend(documents)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
        
        logger.info(f"Processed directory {directory_path}: {len(all_documents)} total chunks")
        return all_documents
    
    def validate_medical_content(self, text: str) -> Dict[str, Any]:
        """
        Validate if the extracted text contains medical content.
        
        Args:
            text: Extracted text to validate
            
        Returns:
            Validation results
        """
        medical_keywords = [
            'diagnosis', 'treatment', 'medication', 'symptom', 'patient',
            'clinical', 'medical', 'drug', 'therapy', 'disease', 'condition',
            'prescription', 'dosage', 'side effect', 'contraindication'
        ]
        
        text_lower = text.lower()
        found_keywords = [keyword for keyword in medical_keywords if keyword in text_lower]
        
        return {
            "is_medical": len(found_keywords) > 0,
            "medical_keywords_found": found_keywords,
            "keyword_count": len(found_keywords),
            "text_length": len(text)
        } 